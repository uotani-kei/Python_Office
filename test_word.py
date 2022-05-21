from docx import Document

document = Document()

document.add_heading('簡単なWordドキュメントのタイトル', 0)
document.add_paragraph('簡単なWordドキュメントのテキスト')

document.save('sample.docx')